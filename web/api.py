"""科研 Agent 系统 Web API。

启动方式：
    python -m web.api
    # 或：uvicorn web.api:app --host 0.0.0.0 --port 8000

提供 REST 接口供前端单页应用调用，所有项目状态存内存（Demo 用，不持久化）。
Pipeline 在独立线程中异步执行，人工节点通过 HumanCallbackBridge 桥接到 REST 接口。
"""
from __future__ import annotations

import os
import shutil
import sys
import threading
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional
from typing import Any, Optional

# 确保项目根在 sys.path（python -m web.api 已保证，但直接运行脚本时需补充）
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

# 启动前加载 .env
from runtime.cli import _load_env  # noqa: E402

_load_env()


def _guard_single_instance() -> None:
    """启动时检测端口冲突，防止多 uvicorn 实例分裂内存状态。

    多实例会导致 _PROJECTS / _BRIDGE（进程内存单例）互相不可见：
    人工节点提交失败、项目进度丢失。本函数在模块 import 时探测端口。

    兼容性：
    - python -m web.api：import 时端口尚未被自己绑定，能准确识别旧实例。
    - uvicorn web.api:app / --reload：worker 子进程 import 本模块时端口尚未被自己
      绑定，不会误判；若检测到旧实例则报错退出（reload supervisor 会反复重启，
      必须先清理旧实例）。
    - 测试/嵌入式场景：设 SRA_WEB_SKIP_GUARD=1 跳过。
    """
    if os.environ.get("SRA_WEB_SKIP_GUARD", "").lower() in ("1", "true", "yes"):
        return
    import socket

    port = int(os.environ.get("SRA_WEB_PORT", "8001"))
    probe = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    try:
        probe.bind(("0.0.0.0", port))
    except OSError:
        print(
            f"[FATAL] 端口 {port} 已被占用，检测到另一个 uvicorn 实例在运行。\n"
            "多实例会导致项目状态互相不可见（人工节点提交失败、进度丢失）。\n"
            "请先停止旧实例再启动：\n"
            f"  netstat -ano | findstr :{port}\n"
            "  taskkill /PID <pid> /F",
            file=sys.stderr,
        )
        sys.exit(1)
    finally:
        probe.close()


_guard_single_instance()

from fastapi import FastAPI, HTTPException  # noqa: E402
from fastapi.middleware.cors import CORSMiddleware  # noqa: E402
from fastapi.responses import FileResponse  # noqa: E402
from fastapi import FastAPI, HTTPException, UploadFile, File, Form  # noqa: E402
from fastapi.middleware.cors import CORSMiddleware  # noqa: E402
from fastapi.responses import FileResponse, Response  # noqa: E402
from fastapi.staticfiles import StaticFiles  # noqa: E402
from pydantic import BaseModel  # noqa: E402

from core.config import get_config  # noqa: E402
from core.knowledge import KnowledgeStore  # noqa: E402
from core.orchestration.node import HumanRequest, HumanResponse  # noqa: E402
from core.state.lifecycle import LifecycleStage  # noqa: E402
from core.state.session import ProjectSession  # noqa: E402
from core.tools.journal_quality import enrich_paper_quality, build_pdf_url  # noqa: E402
from core.tools.doi_resolve import find_open_access_pdf, resolve_doi_by_title  # noqa: E402


# PDF 解析结果内存缓存（key = doi 或 title 前 120 字符 → pdf_url）
# 避免 list_papers 每刷新一次页面就重复发起 Crossref/Unpaywall 网络请求
_pdf_resolve_cache: dict[str, str] = {}
_pdf_resolve_lock = threading.Lock()
from core.tools.doi_resolve import find_open_access_pdf, resolve_doi_by_title  # noqa: E402
from runtime.pipeline import Pipeline, PipelineResult  # noqa: E402
from stages.common import RESEARCH_SEARCH_PREFS  # noqa: E402


# ===== 请求/响应模型 =====


class CreateProjectRequest(BaseModel):
    topic: str


class TopicDiscoveryRequest(BaseModel):
    """方向推荐请求。"""
    interest: str
class RunProjectRequest(BaseModel):
    force_writing: bool = False


class NoteRequest(BaseModel):
    text: str


class HumanResponseRequest(BaseModel):
    action: str = "continue"  # continue / abort / rollback
    text: Optional[str] = None
    selected_option: Optional[str] = None
    # 检索偏好（可选）：年份区间 + 期刊关键词，在确认检索方向节点提交
    search_prefs: Optional[dict] = None


# ===== 人工节点桥接 =====


class HumanCallbackBridge:
    """跨线程的人工节点桥接。

    Pipeline 在工作线程中执行，遇到 HumanNode 时调用 callback 阻塞等待；
    主线程的 REST 接口接收用户响应后唤醒 callback，让 Pipeline 继续。
    """

    def __init__(self) -> None:
        self._lock = threading.Lock()
        # project_id -> 等待中的人工请求描述（dict 形式，便于序列化）
        self._pending: dict[str, dict[str, Any]] = {}
        # project_id -> threading.Event，用于阻塞 callback
        self._events: dict[str, threading.Event] = {}
        # project_id -> HumanResponse，由 REST 接口注入
        self._responses: dict[str, HumanResponse] = {}

    def make_callback(
        self,
        project_id: str,
        on_pending: Optional[Callable[[], None]] = None,
        on_resume: Optional[Callable[[], None]] = None,
    ):
        def _cb(req: HumanRequest) -> HumanResponse:
            event = threading.Event()
            payload = {
                "prompt": req.prompt,
                "options": list(req.options) if req.options else [],
                "allow_free_text": req.allow_free_text,
                "context": dict(req.context) if req.context else {},
                "appeared_at": datetime.now().isoformat(),
            }
            with self._lock:
                self._pending[project_id] = payload
                self._events[project_id] = event
            # 通知外部（工作线程）项目进入人工等待状态
            if on_pending is not None:
                on_pending()
            # 阻塞直到 REST 接口提交响应或超时
            triggered = event.wait(timeout=3600)
            # 收到响应/超时后恢复 running
            if on_resume is not None:
                on_resume()
            with self._lock:
                resp = self._responses.pop(project_id, None)
                self._pending.pop(project_id, None)
                self._events.pop(project_id, None)
            if not triggered or resp is None:
                return HumanResponse(action="abort")
            return resp

        return _cb

    def submit(self, project_id: str, response: HumanResponse) -> bool:
        """注入响应并唤醒 callback。返回是否成功（存在等待中的请求）。"""
        with self._lock:
            event = self._events.get(project_id)
            if event is None:
                return False
            self._responses[project_id] = response
        event.set()
        return True

    def get_pending(self, project_id: str) -> Optional[dict[str, Any]]:
        with self._lock:
            req = self._pending.get(project_id)
            return dict(req) if req else None

    def clear(self, project_id: str) -> None:
        with self._lock:
            self._pending.pop(project_id, None)
            self._responses.pop(project_id, None)
            event = self._events.pop(project_id, None)
        if event is not None:
            event.set()


# ===== 项目状态 =====


@dataclass
class ProjectState:
    project_id: str
    topic: str
    created_at: str
    status: str = "created"  # created / running / pending_human / completed / failed / aborted
    summary: str = ""
    last_result: Optional[PipelineResult] = None
    error: Optional[str] = None
    thread: Optional[threading.Thread] = None
    node_history: list[dict] = field(default_factory=list)
    # 当前正在执行的节点 + 下一步候选（长任务实时进度提示）
    current_node: Optional[dict] = None
    next_nodes: list[str] = field(default_factory=list)
    notes: list[dict] = field(default_factory=list)
    recommendation: str = ""
    # 路线 A：构效关系发现产出（从 discovery 子图 context 收集）
    discovery: dict = field(default_factory=dict)
    # 方向推荐产出（从 topic_discovery 子图收集）
    topic_discovery: dict = field(default_factory=dict)
    # 当前运行模式：pipeline / discovery / topic_discovery
    # 当前运行模式：pipeline / discovery
    run_mode: str = ""
    # 用户检索偏好（论文抓取范围）：{year_min, year_max, venue_hint}
    search_prefs: dict = field(default_factory=dict)


# ===== 全局状态 =====

_CONFIG = get_config()
_PROJECTS: dict[str, ProjectState] = {}
_BRIDGE = HumanCallbackBridge()
_LOCK = threading.Lock()


def _recover_project_topic(pid: str) -> str:
    """服务重启后从磁盘恢复项目主题。

    优先级：
    1. KV `research.topic`（pipeline start/resume 时持久化，最准确）
    2. 调研报告 cross_validation_report 里的子问题（subquery 字段，带语义）
    3. 兜底空字符串（前端提示重新确认主题）
    """
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(pid))
        topic = store.get_kv("research.topic", "")
        if topic:
            return str(topic)
        # 兜底：从交叉验证报告提取信息最丰富的子问题（去掉疑问后缀，保留主干）
        report = store.get_kv("cross_validation_report", {}) or {}
        for key in ("gaps", "consensus", "conflicts"):
            items = report.get(key) or []
            for it in items:
                sq = (it or {}).get("subquery") if isinstance(it, dict) else None
                if sq and str(sq).strip():
                    # 子问题语义最接近主题，直接作为恢复主题（保留完整可读）
                    return str(sq).strip()[:60]
    except Exception:  # noqa: BLE001
        pass
    return ""


def _scan_existing_projects() -> None:
    """启动时从磁盘扫描已存在项目并恢复内存状态。

    项目数据（knowledge.db / snapshots）持久化在 projects/ 目录，
    但 ProjectState 是进程内存单例，服务重启即丢失。此函数让重启后
    旧项目仍可通过 Web 访问（论文/材料/证据等数据按 project_id 读取）。
    topic 不持久化（ProjectSession 快照不含），恢复项目显示 topic 为空，
    重新运行前需在 UI 确认主题。
    """
    projects_dir = _CONFIG.paths.projects
    if not projects_dir.exists():
        return
    for pdir in sorted(projects_dir.iterdir()):
        if not pdir.is_dir() or pdir.name in _PROJECTS:
            continue
        if not (pdir / "knowledge.db").exists():
            continue
        pid = pdir.name
        # created_at 从目录名解析：proj_YYYYMMDD_HHMMSS_xxxx
        created_at = datetime.utcnow().isoformat()
        try:
            _, ts = pid.split("_", 1)
            created_at = datetime.strptime(ts[:13], "%Y%m%d_%H%M%S").isoformat()
        except Exception:
            pass
        # 从 session 快照推断运行状态
        status, summary = "created", ""
        try:
            session = ProjectSession.load(pid, _CONFIG.paths)
            cur = session.current_stage().value
            done = [s for s in LifecycleStage.ordered()
                    if session.status_of(s).value in ("completed", "blocked", "failed")]
            if done or cur != LifecycleStage.RESEARCH.value:
                status = "completed"
                summary = f"服务重启后自动恢复（最近阶段: {cur}）"
            else:
                summary = "服务重启后自动恢复（数据可查看，重新运行需确认主题）"
        except Exception:
            status, summary = "created", "服务重启后自动恢复"
        _PROJECTS[pid] = ProjectState(
            project_id=pid,
            topic=_recover_project_topic(pid),
            created_at=created_at,
            status=status,
            summary=summary,
        )


_scan_existing_projects()


# ===== Pipeline 工作线程 =====


def _set_state_status(state: ProjectState, status: str) -> None:
    """线程安全更新项目状态。"""
    with _LOCK:
        state.status = status


def _make_progress_callback(
    state: ProjectState,
) -> tuple[Callable[[list[dict]], None], Callable[[str, list[str]], None]]:
    """构造节点级实时进度回调对（完成回调 + 开始回调）。

    完成回调：每完成一个节点就把最新节点历史写入 state.node_history。
    开始回调：每开始一个节点就把当前节点 ID 与下一步候选写入 state。

    这是解决“页面长时间无反应”的关键：research 阶段真实运行需 15-25 分钟，
    若只在 pipeline 结束后才写 node_history，前端会一直空白。此回调让
    /status 轮询能实时看到节点进度，并展示「正在执行 / 下一步」避免干等。

    注意：只写 node_history / current_node，不改 state.status。status 由
    线程函数与 on_pending/on_resume（pending_human ↔ running）统一管理，
    避免回调在人工等待期间误覆盖状态。
    """
    def _on_progress(history: list[dict]) -> None:
        with _LOCK:
            state.node_history = history

    def _on_node_started(node_id: str, next_nodes: list[str]) -> None:
        with _LOCK:
            state.current_node = {
                "node_id": node_id,
                "next_nodes": next_nodes,
                "started_at": datetime.utcnow().isoformat(),
            }
            state.next_nodes = list(next_nodes)

    return _on_progress, _on_node_started


# ===== Pipeline 工作线程 =====


def _run_pipeline_thread(project_id: str, topic: str, resume: bool, force_writing: bool = False) -> None:
    """工作线程函数：执行 Pipeline 并更新项目状态。"""
    state = _PROJECTS.get(project_id)
    if state is None:
        return
    try:
        _set_state_status(state, "running")
        pipeline = Pipeline(config=_CONFIG)
        # dry_run 模式下也使用桥接回调，让用户能介入；非 dry_run 同样使用桥接
        # 人工等待期间将 status 置为 pending_human，前端可明确感知
        human_cb = _BRIDGE.make_callback(
            project_id,
            on_pending=lambda: _set_state_status(state, "pending_human"),
            on_resume=lambda: _set_state_status(state, "running"),
        )
        _progress_cbs = _make_progress_callback(state)
        state.status = "running"
        pipeline = Pipeline(config=_CONFIG)
        # dry_run 模式下也使用桥接回调，让用户能介入；非 dry_run 同样使用桥接
        human_cb = _BRIDGE.make_callback(project_id)
        result: PipelineResult = pipeline.run_pipeline(
            project_id=project_id,
            topic=topic,
            human_callback=human_cb,
            resume=resume,
            on_progress=_progress_cbs[0],
            on_node_started=_progress_cbs[1],
            force_writing=force_writing,
            initial_ctx={RESEARCH_SEARCH_PREFS: state.search_prefs} if state.search_prefs else None,
        )
        state.last_result = result
        state.status = result.status
        state.summary = result.summary
        state.recommendation = result.recommendation
        state.node_history = result.node_history or []
        state.current_node = None
        state.next_nodes = []
    except Exception as e:  # noqa: BLE001
        state.status = "failed"
        state.error = f"{type(e).__name__}: {e}"
        state.summary = f"Pipeline 执行异常: {e}"


def _run_discovery_thread(project_id: str, topic: str, resume: bool) -> None:
    """工作线程函数：执行构效关系发现（路线 A）并更新项目状态。"""
    state = _PROJECTS.get(project_id)
    if state is None:
        return
    try:
        _set_state_status(state, "running")
        state.run_mode = "discovery"
        pipeline = Pipeline(config=_CONFIG)
        human_cb = _BRIDGE.make_callback(
            project_id,
            on_pending=lambda: _set_state_status(state, "pending_human"),
            on_resume=lambda: _set_state_status(state, "running"),
        )
        _progress_cbs = _make_progress_callback(state)
        state.status = "running"
        state.run_mode = "discovery"
        pipeline = Pipeline(config=_CONFIG)
        human_cb = _BRIDGE.make_callback(project_id)
        result: PipelineResult = pipeline.run_discovery(
            project_id=project_id,
            topic=topic,
            human_callback=human_cb,
            resume=resume,
            on_progress=_progress_cbs[0],
            on_node_started=_progress_cbs[1],
        )
        state.last_result = result
        state.status = result.status
        state.summary = result.summary
        state.recommendation = result.recommendation
        state.node_history = result.node_history or []
        state.current_node = None
        state.next_nodes = []
        # 从 node_history 提取 discovery 产出摘要
        state.discovery = _extract_discovery_summary(
            result.node_history, result.extra.get("hypotheses") or []
        )
        # 从 node_history 提取 discovery 产出摘要
        state.discovery = _extract_discovery_summary(result.node_history)
    except Exception as e:  # noqa: BLE001
        state.status = "failed"
        state.error = f"{type(e).__name__}: {e}"
        state.summary = f"Discovery 执行异常: {e}"


def _extract_discovery_summary(node_history: list[dict], hypotheses: list[dict] | None = None) -> dict:
    """从节点历史提取 discovery 产出摘要。"""
    summary = {
        "hypotheses": 0, "candidates": 0, "relationships": 0, "novel": 0,
        "nodes": [], "hypothesis_list": [],
    }
def _extract_discovery_summary(node_history: list[dict]) -> dict:
    """从节点历史提取 discovery 产出摘要。"""
    summary = {"hypotheses": 0, "candidates": 0, "relationships": 0, "novel": 0, "nodes": []}
    for h in node_history or []:
        node_id = h.get("node_id", "")
        if node_id in ("hypothesis_seed", "search_space", "llm_guided_search",
                        "discovery_validate", "discovery_report"):
            summary["nodes"].append({
                "node_id": node_id,
                "status": h.get("status"),
                "summary": h.get("summary", ""),
            })
            if node_id == "hypothesis_seed" and "个候选构效关系假设" in h.get("summary", ""):
                try:
                    summary["hypotheses"] = int(h["summary"].split("生成")[1].split("个")[0])
                except (IndexError, ValueError):
                    pass
            if node_id == "discovery_validate" and "条验证发现" in h.get("summary", ""):
                try:
                    summary["relationships"] = int(
                        h["summary"].split("验证")[1].split("条")[0])
                    if "条 novel" in h["summary"]:
                        summary["novel"] = int(
                            h["summary"].split("其中 ")[1].split(" 条")[0])
                except (IndexError, ValueError):
                    pass
    # 假设列表（含三维可验证性评分），按综合分降序，供 Web 排序展示
    h_list: list[dict] = []
    for hyp in hypotheses or []:
        if not isinstance(hyp, dict) or not hyp.get("hypothesis"):
            continue
        def _f(key: str) -> float:
            try:
                return float(hyp.get(key, 0.0))
            except (TypeError, ValueError):
                return 0.0
        n_, f_, g_ = _f("novelty_score"), _f("feasibility_score"), _f("gap_relevance_score")
        h_list.append({
            "hypothesis": hyp.get("hypothesis", ""),
            "variables": hyp.get("variables", []) or [],
            "target_property": hyp.get("target_property", ""),
            "rationale": hyp.get("rationale", ""),
            "gap_ref": hyp.get("gap_ref", ""),
            "novelty_score": round(n_, 2),
            "feasibility_score": round(f_, 2),
            "gap_relevance_score": round(g_, 2),
            "overall_score": round(0.4 * n_ + 0.3 * f_ + 0.3 * g_, 2),
        })
    h_list.sort(key=lambda x: x["overall_score"], reverse=True)
    summary["hypothesis_list"] = h_list
    return summary


def _run_topic_discovery_thread(project_id: str, interest: str, resume: bool) -> None:
    """工作线程函数：执行方向推荐并更新项目状态。"""
    state = _PROJECTS.get(project_id)
    if state is None:
        return
    try:
        _set_state_status(state, "running")
        state.run_mode = "topic_discovery"
        pipeline = Pipeline(config=_CONFIG)
        human_cb = _BRIDGE.make_callback(
            project_id,
            on_pending=lambda: _set_state_status(state, "pending_human"),
            on_resume=lambda: _set_state_status(state, "running"),
        )
        def _publish_recommendations(recs: list[dict]) -> None:
            """推荐就绪时提前写入 state，供 pending_human 状态下前端展示卡片。"""
            with _LOCK:
                state.topic_discovery = {
                    **state.topic_discovery,
                    "recommendations": recs,
                    "interest": interest,
                }

        _progress_cbs = _make_progress_callback(state)
        result: PipelineResult = pipeline.run_topic_discovery(
            project_id=project_id,
            interest=interest,
            human_callback=human_cb,
            auto_research=False,
            resume=resume,
            on_recommendations=_publish_recommendations,
            on_progress=_progress_cbs[0],
            on_node_started=_progress_cbs[1],
        )
        state.last_result = result
        state.status = result.status
        state.summary = result.summary
        state.recommendation = result.recommendation
        state.node_history = result.node_history or []
        state.current_node = None
        state.next_nodes = []
        # 存储方向推荐完整数据（推荐列表 + 选择结果）
        state.topic_discovery = result.extra or {}
        # 若用户已选定主题，自动把所选主题写入 state.topic，并启动 research pipeline
        selected = (result.extra or {}).get("selected_topic", "") or ""
        if result.status == "completed" and selected.strip():
            with _LOCK:
                state.topic = selected.strip()
                state.run_mode = "pipeline"
            # 在新线程中异步启动 research pipeline，避免阻塞 topic_discovery 线程
            resume2 = True  # 复用已有 project_dir 中的 snapshots
            pipeline_thread = threading.Thread(
                target=_run_pipeline_thread,
                args=(project_id, selected.strip(), resume2, False),
                daemon=True,
            )
            pipeline_thread.start()
    except Exception as e:  # noqa: BLE001
        state.status = "failed"
        state.error = f"{type(e).__name__}: {e}"
        state.summary = f"TopicDiscovery 执行异常: {e}"


def _extract_topic_discovery_summary(node_history: list[dict]) -> dict:
    """从节点历史提取方向推荐产出摘要。"""
    summary = {
        "trends_fetched": False,
        "emerging_count": 0,
        "stable_count": 0,
        "saturated_count": 0,
        "recommendations_count": 0,
        "selected_topic": "",
        "nodes": [],
    }
    for h in node_history or []:
        node_id = h.get("node_id", "")
        if node_id in ("trend_fetch", "trend_analysis", "topic_recommend", "topic_select"):
            summary["nodes"].append({
                "node_id": node_id,
                "status": h.get("status"),
                "summary": h.get("summary", ""),
            })
            if node_id == "trend_fetch":
                summary["trends_fetched"] = True
            if node_id == "trend_analysis":
                # 从 summary 提取数量
                s = h.get("summary", "")
                try:
                    if "个新兴方向" in s:
                        summary["emerging_count"] = int(s.split("个新兴方向")[0].split("：")[-1])
                    if "个稳定方向" in s:
                        summary["stable_count"] = int(s.split("个稳定方向")[0].split("，")[-1])
                    if "个饱和方向" in s:
                        summary["saturated_count"] = int(s.split("个饱和方向")[0].split("，")[-1])
                except (IndexError, ValueError):
                    pass
            if node_id == "topic_recommend":
                s = h.get("summary", "")
                try:
                    if "个推荐主题" in s:
                        summary["recommendations_count"] = int(s.split("生成")[1].split("个")[0])
                except (IndexError, ValueError):
                    pass
            if node_id == "topic_select":
                s = h.get("summary", "")
                # 用户选择的主题会通过 context 写入，这里只是标记
                summary["selected_topic"] = ""
    return summary


# ===== FastAPI 应用 =====


app = FastAPI(title="科研 Agent 系统", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


STATIC_DIR = Path(__file__).resolve().parent / "static"


# ===== 接口实现 =====


@app.post("/api/projects")
def create_project(req: CreateProjectRequest) -> dict:
    """创建新项目，返回 project_id。"""
    topic = (req.topic or "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="topic 不能为空")
    project_id = f"proj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
    state = ProjectState(
        project_id=project_id,
        topic=topic,
        created_at=datetime.utcnow().isoformat(),
    )
    with _LOCK:
        _PROJECTS[project_id] = state
    return {"project_id": project_id, "topic": topic, "created_at": state.created_at}


@app.get("/api/projects")
def list_projects() -> dict:
    """列出所有项目（供前端刷新后恢复）。"""
    with _LOCK:
        items = [
            {
                "project_id": s.project_id,
                "topic": s.topic,
                "created_at": s.created_at,
                "status": s.status,
                "summary": s.summary,
            }
            for s in _PROJECTS.values()
        ]
    # 按创建时间降序
    items.sort(key=lambda x: x["created_at"], reverse=True)
    return {"projects": items}


@app.delete("/api/projects/{project_id}")
def delete_project(project_id: str) -> dict:
    """删除项目：清理内存中的 ProjectState，并删除磁盘上的项目目录（含 knowledge.db）。"""
    # 内存清理
    with _LOCK:
        state = _PROJECTS.pop(project_id, None)
        if state is not None and state.thread and state.thread.is_alive():
            # 线程仍在运行：不允许删除，避免留下孤立线程
            _PROJECTS[project_id] = state
            raise HTTPException(
                status_code=409,
                detail="项目正在运行中，请先中止 Pipeline 再删除",
            )
    # 磁盘清理：项目目录（knowledge.db / vectors / snapshots / artifacts / uploads / experiments）
    project_dir = _CONFIG.paths.project_dir(project_id)
    if project_dir.exists() and project_dir.resolve().is_relative_to(_CONFIG.paths.projects.resolve()):
        shutil.rmtree(project_dir, ignore_errors=True)
    return {"project_id": project_id, "deleted": True}


@app.get("/api/projects/{project_id}/status")
def get_status(project_id: str) -> dict:
    state = _require_project(project_id)
    # 当前阶段：从 session 读取（若已启动）
    current_stage = ""
    stage_statuses: dict[str, str] = {}
    try:
        session = ProjectSession.load(project_id, _CONFIG.paths)
        current_stage = session.current_stage().value
        for stage in LifecycleStage.ordered():
            stage_statuses[stage.value] = session.status_of(stage).value
    except Exception:  # noqa: BLE001
        pass

    # 产出物计数
    counts = {"papers": 0, "ideas": 0, "claims": 0, "experiments": 0, "evidence": 0,
              "materials": 0, "properties": 0, "synthesis": 0, "gaps": 0}
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        counts["papers"] = len(store.list_papers())
        counts["ideas"] = len(store.list_ideas())
        counts["claims"] = len(store.list_claims())
        counts["experiments"] = len(store.list_experiments())
        counts["evidence"] = store.evidence_stats()["total"]
        mstats = store.material_stats()
        counts["materials"] = mstats["total_materials"]
        counts["properties"] = mstats["total_properties"]
        counts["synthesis"] = mstats["total_synthesis"]
        counts["gaps"] = store.gap_stats()["total"]
    except Exception:  # noqa: BLE001
        pass

    pending = _BRIDGE.get_pending(project_id)

    return {
        "project_id": project_id,
        "topic": state.topic,
        "created_at": state.created_at,
        "status": state.status,
        "summary": state.summary,
        "recommendation": state.recommendation,
        "error": state.error,
        "current_stage": current_stage,
        "stage_statuses": stage_statuses,
        "node_history": state.node_history,
        "current_node": state.current_node,
        "next_nodes": state.next_nodes,
        "counts": counts,
        "pending_human": pending,
        "search_prefs": state.search_prefs,
        "topic_discovery": state.topic_discovery,
        "counts": counts,
        "pending_human": pending,
    }


@app.post("/api/projects/{project_id}/run")
def run_project(project_id: str, req: Optional[RunProjectRequest] = None) -> dict:
    """启动/继续 pipeline（异步执行，立即返回）。

    可选 body: {"force_writing": true} —— 实验失败后强制进入论文写作阶段。
    """
    state = _require_project(project_id)
    if state.status == "running":
        raise HTTPException(status_code=409, detail="pipeline 正在运行中")
    if state.status == "pending_human":
        raise HTTPException(
            status_code=409,
            detail="当前等待人工响应，请先提交 human-response 或中止",
        )
    if not state.topic.strip():
        raise HTTPException(
            status_code=400,
            detail="项目主题为空（服务重启后恢复的项目不持久化 topic），请重新创建项目或补充主题",
        )
    resume = state.status not in ("created",)
    force_writing = (req.force_writing if req else False)
    thread = threading.Thread(
        target=_run_pipeline_thread,
        args=(project_id, state.topic, resume, force_writing),
        daemon=True,
    )
    state.thread = thread
    state.error = None
    thread.start()
    msg = "pipeline 已启动"
    if force_writing:
        msg = "pipeline 已启动（强制写作模式：绕过实验成败判断）"
    return {"project_id": project_id, "message": msg, "resumed": resume, "force_writing": force_writing}


@app.post("/api/projects/{project_id}/run-discovery")
def run_discovery(project_id: str) -> dict:
    """启动构效关系发现（路线 A）：research → discovery 子图。

    异步执行，立即返回。结果通过 /status 与 /discoveries 接口查询。
    """
    state = _require_project(project_id)
    if state.status == "running":
        raise HTTPException(status_code=409, detail="任务正在运行中")
    if state.status == "pending_human":
        raise HTTPException(
            status_code=409,
            detail="当前等待人工响应，请先提交 human-response 或中止",
        )
    if not state.topic.strip():
        raise HTTPException(
            status_code=400,
            detail="项目主题为空（服务重启后恢复的项目不持久化 topic），请重新创建项目或补充主题",
        )
    resume = state.status not in ("created",)
    thread = threading.Thread(
        target=_run_discovery_thread,
        args=(project_id, state.topic, resume),
        daemon=True,
    )
    state.thread = thread
    state.error = None
    thread.start()
    return {"project_id": project_id, "message": "discovery 已启动", "resumed": resume}


@app.get("/api/projects/{project_id}/discoveries")
def get_discoveries(project_id: str) -> dict:
    """获取构效关系发现产出（路线 A）。

    返回内容包括：
    - relationships：构效关系 Claim 列表
    - discovery_summary：发现汇总（含 novel 计数）
    - reliability_scores：客观可信度评分（5 维度 + 风险标签）
    - expert_assistance：专家辅助包（最近邻工艺 + 类似材料 + DFT + 实验 protocol）
    - gap_scores：上游 Research Gap 质量评分（4 维度 + 综合分）
    - run_mode：运行模式
    """
    state = _require_project(project_id)
    # 从 KnowledgeStore 读取 discovery 阶段产出的 Claim（构效关系发现）
    relationships: list[dict] = []
    reliability_scores: list[dict] = []
    expert_assistances: list[dict] = []
    gap_scores: list[dict] = []
    gap_summary: dict = {}
    rel_summary: dict = {}
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        # Claim
        for c in store.list_claims():
            if c.source_stage == "discovery":
                relationships.append({
                    "claim_id": c.claim_id,
                    "statement": c.statement,
                    "status": c.status.value,
                    "evidence_refs": c.evidence_refs,
                    "created_at": c.created_at.isoformat() if c.created_at else None,
                })
        # 客观可信度评分（路线 A 客观指标：5 维度）
        rel_kv = store.get_kv("discovery_reliability_scores") or {}
        reliability_scores = rel_kv.get("scores", []) if isinstance(rel_kv, dict) else []
        rel_summary = rel_kv.get("summary", {}) if isinstance(rel_kv, dict) else {}
        # 专家辅助包（让材料专家感到"对我有用"）
        expert_assistances = store.get_kv("discovery_expert_assistance") or []
        # 上游 Gap 质量评分
        gap_kv = store.get_kv("research_gap_scores") or {}
        gap_scores = gap_kv.get("scores", []) if isinstance(gap_kv, dict) else []
        gap_summary = gap_kv.get("summary", {}) if isinstance(gap_kv, dict) else {}
        # 兼容旧数据：评分里没有 statement/paper_titles 时，用 research_gaps 表回填
        # （gap_id → statement + evidence 论文标题），保证前端始终可读
        if gap_scores:
            try:
                gaps_in_store = store.list_research_gaps()
                gap_meta = {
                    g.gap_id: g for g in gaps_in_store
                }
                enriched = []
                for s in gap_scores:
                    if s.get("statement") or s.get("paper_titles"):
                        enriched.append(s)
                        continue
                    s2 = dict(s)
                    g = gap_meta.get(s.get("gap_id", ""))
                    if g:
                        s2["statement"] = g.statement or ""
                        evs = g.evidence or []
                        s2["evidence"] = evs
                        s2["paper_titles"] = [
                            e.get("title", "") for e in evs if e.get("title")
                        ][:5]
                        s2["paper_ids"] = [
                            e.get("paper_id", "") for e in evs if e.get("paper_id")
                        ][:5]
                    enriched.append(s2)
                gap_scores = enriched
            except Exception:  # noqa: BLE001
                pass
    except Exception as e:  # noqa: BLE001
        pass
    return {
        "project_id": project_id,
        "discovery_summary": state.discovery,
        "relationships": relationships,
        "reliability_scores": reliability_scores,
        "reliability_summary": rel_summary,
        "expert_assistances": expert_assistances,
        "gap_scores": gap_scores,
        "gap_summary": gap_summary,
        "run_mode": state.run_mode,
    }


@app.post("/api/projects/{project_id}/run-topic-discovery")
def run_topic_discovery(project_id: str, req: TopicDiscoveryRequest) -> dict:
    """启动方向推荐（topic_discovery 子图）。

    在用户给定主题之前，主动分析领域趋势、推荐研究主题。
    异步执行，推荐结果通过 /recommendations 接口查询，
    用户选择通过 /human-response 接口提交。
    """
    state = _require_project(project_id)
    interest = (req.interest or "").strip()
    if not interest:
        raise HTTPException(status_code=400, detail="interest 不能为空")
    if state.status == "running":
        raise HTTPException(status_code=409, detail="任务正在运行中")
    if state.status == "pending_human":
        raise HTTPException(
            status_code=409,
            detail="当前等待人工响应，请先提交 human-response 或中止",
        )
    resume = state.status not in ("created",)
    thread = threading.Thread(
        target=_run_topic_discovery_thread,
        args=(project_id, interest, resume),
        daemon=True,
    )
    state.thread = thread
    state.error = None
    state.topic = interest  # 更新 topic 为用户输入的研究兴趣
    thread.start()
    return {"project_id": project_id, "message": "topic_discovery 已启动", "interest": interest}


@app.get("/api/projects/{project_id}/recommendations")
def get_recommendations(project_id: str) -> dict:
    """获取方向推荐产出（推荐主题列表 + 趋势分析摘要）。"""
    state = _require_project(project_id)
    return {
        "project_id": project_id,
        "topic_discovery_summary": state.topic_discovery,
        "run_mode": state.run_mode,
    }


@app.get("/api/projects/{project_id}/papers")
def list_papers(project_id: str) -> dict:
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        papers = store.list_papers()
    except Exception as e:  # noqa: BLE001
        return {"papers": [], "error": f"读取失败: {e}"}

    result = []
    for p in papers:
        # 即时补充期刊质量字段（对旧数据兜底：入库时未填充的在此补上）
        # 注意：这里只做本地构造（arxiv_id → pdf、已有 doi → doi.org），
        # 不发起 OpenAlex/Crossref 网络请求（避免接口阻塞超时）。
        # DOI 反查 + OA PDF 查找由后台补全脚本（_backfill_pdf_quality.py）持久化写回 DB。
        meta = p.metadata or {}
        if "impact_factor" not in meta:
            # 旧论文 metadata 里没有质量字段，即时补充
            tmp = {
                "title": p.title or "",
                "authors": p.authors or [],
                "year": p.year,
                "venue": p.venue or "",
                "arxiv_id": p.arxiv_id or "",
                "doi": p.doi or meta.get("doi", ""),
                "pdf_url": meta.get("pdf_url", ""),
            }
            try:
                enrich_paper_quality(tmp)
                # PDF 链接：已有直链 / arxiv_id 直接构造；有 DOI 用 doi.org 兜底
                # （读进程内缓存，命中则直接使用；未命中只做本地构造，不联网）
                if not (tmp.get("pdf_url") or "").strip():
                    tmp["pdf_url"] = build_pdf_url(tmp)
                    if not tmp["pdf_url"] and (tmp.get("doi") or "").strip():
                        tmp["pdf_url"] = f"https://doi.org/{tmp['doi']}"
            except Exception:  # noqa: BLE001
                pass
            meta = {**meta, "impact_factor": tmp.get("impact_factor", 0.0),
                    "cas_zone": tmp.get("cas_zone", ""),
                    "cas_subcategory": tmp.get("cas_subcategory", ""),
                    "is_top_journal": tmp.get("is_top_journal", False),
                    "pdf_url": tmp.get("pdf_url", ""),
                    "doi": tmp.get("doi", p.doi or "")}

        result.append({
            "paper_id": p.paper_id,
            "title": p.title,
            "authors": p.authors,
            "year": p.year,
            "venue": p.venue,
            "arxiv_id": p.arxiv_id,
            "abstract": p.abstract,
            "url": p.url or (f"https://arxiv.org/abs/{p.arxiv_id}" if p.arxiv_id else None),
            "doi": p.doi or meta.get("doi", ""),
            "doi_url": (f"https://doi.org/{p.doi or meta.get('doi', '')}" if (p.doi or meta.get("doi", "")) else None),
            "pdf_path": p.pdf_path,
            "created_at": p.created_at.isoformat() if p.created_at else None,
            "source_stage": p.source_stage,
            # 检索证据链字段（审计溯源：来源 / doc_id / 证据分 / 命中子问题）
            "source": meta.get("source", ""),
            "source_subquery": meta.get("source_subquery", ""),
            "doc_id": meta.get("doc_id", ""),
            "offset": meta.get("offset", 0),
            "evidence_score": meta.get("evidence_score", 0.0),
            "relevance_score": meta.get("relevance_score", 0.0),
            "relevance_reason": meta.get("relevance_reason", ""),
            # 期刊质量指标（journal_quality 模块填充）
            "impact_factor": meta.get("impact_factor", 0.0),
            "cas_zone": meta.get("cas_zone", ""),
            "cas_subcategory": meta.get("cas_subcategory", ""),
            "is_top_journal": meta.get("is_top_journal", False),
            "pdf_url": meta.get("pdf_url", ""),
        })
    return {"papers": result}


@app.get("/api/projects/{project_id}/materials")
def list_materials(project_id: str) -> dict:
    """获取材料知识库（Task 2：材料-性能-合成三元组）。

    返回按材料聚合的知识：每种材料含其性能指标与合成方法，
    均带来源论文与证据片段（可溯源）。满足赛题「知识抽取结构化」要求。

    数据标准化（core/knowledge/normalize.py）：
    - 性能指标统一映射（ZT/zT/figure_of_merit → 热电优值，带标准符号/单位/类别）
    - 合成方法归类（固相法/熔融法/烧结法/计算模拟…）
    - 材料体系分类（Bi₂Te₃ 基 / PbTe 基 / 钙钛矿 / 方钴矿…）
    """
    from core.knowledge.normalize import (
        categorize_material,
        classify_method,
        normalize_property,
    )

    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        materials = store.list_materials(limit=500)
        stats = store.material_stats()
        # 聚合：材料 → 性能列表 + 合成列表
        properties = store.list_material_properties()
        synthesis = store.list_material_synthesis()
        props_by_mat: dict[str, list[dict]] = {}
        for p in properties:
            norm = normalize_property(p.property_name, p.property_name_cn)
            props_by_mat.setdefault(p.material_id, []).append({
                "property_name": p.property_name,
                "property_name_cn": p.property_name_cn,
                # 标准化字段
                "norm_key": norm["key"],
                "norm_cn": norm["cn"],
                "symbol": norm["symbol"],
                "unit": norm["unit"],
                "category": norm["category"],
                "value": p.value,
                "value_num": p.value_num,
                "condition": p.condition,
                "paper_title": p.paper_title,
                "confidence": p.confidence,
                "source_snippet": p.source_snippet,
            })
        syn_by_mat: dict[str, list[dict]] = {}
        for s in synthesis:
            mth = classify_method(s.method)
            syn_by_mat.setdefault(s.material_id, []).append({
                "method": s.method,
                # 标准化字段
                "method_category": mth["category"],
                "method_label": mth["label"],
                "precursors": s.precursors,
                "temperature": s.temperature,
                "pressure": s.pressure,
                "atmosphere": s.atmosphere,
                "duration": s.duration,
                "steps": s.steps,
                "paper_title": s.paper_title,
                "confidence": s.confidence,
                "source_snippet": s.source_snippet,
            })
        items = []
        prop_cat_counter: dict[str, int] = {}
        method_cat_counter: dict[str, int] = {}
        material_cat_counter: dict[str, int] = {}
        for m in materials:
            mat_cat = categorize_material(m.name, m.formula)
            material_cat_counter[mat_cat] = material_cat_counter.get(mat_cat, 0) + 1
            props = props_by_mat.get(m.material_id, [])
            syns = syn_by_mat.get(m.material_id, [])
            for pr in props:
                prop_cat_counter[pr["category"]] = prop_cat_counter.get(pr["category"], 0) + 1
            for sy in syns:
                method_cat_counter[sy["method_category"]] = method_cat_counter.get(sy["method_category"], 0) + 1
            items.append({
                "material_id": m.material_id,
                "name": m.name,
                "formula": m.formula,
                "category": mat_cat,
                "crystal_structure": m.crystal_structure,
                "space_group": m.space_group,
                "lattice_parameters": m.lattice_parameters,
                "symmetry": m.symmetry,
                "composition": m.composition,
                "paper_id": m.paper_id,
                "paper_title": m.paper_title,
                "source_paper_ids": m.metadata.get("source_paper_ids", []),
                "confidence": m.confidence,
                "source_snippet": m.source_snippet,
                "properties": props,
                "synthesis": syns,
            })
    except Exception as e:  # noqa: BLE001
        return {"materials": [], "stats": {}, "error": f"读取失败: {e}"}
    return {
        "materials": items,
        "stats": stats,
        "aggregation": {
            "property_categories": prop_cat_counter,
            "method_categories": method_cat_counter,
            "material_categories": material_cat_counter,
        },
    }


@app.get("/api/projects/{project_id}/gaps")
def list_research_gaps(project_id: str) -> dict:
    """获取研究缺口清单（Task 3：Research Gap 识别）。

    由 ResearchGapIdentifyAgent 在 cross_validate 之后生成并落库，
    每条 Gap 含类型（矛盾结论/未被探索方向/缺失知识连接）、
    证据链（可溯源 paper_id + snippet）、可操作性、优先级，
    满足赛题「文献溯源完整性」要求，供前端「研究缺口」页展示。
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        gaps = store.list_research_gaps(limit=200)
        stats = store.gap_stats()
    except Exception as e:  # noqa: BLE001
        return {"gaps": [], "stats": {"total": 0, "by_type": {}},
                "error": f"读取失败: {e}"}
    return {
        "gaps": [
            {
                "gap_id": g.gap_id,
                "gap_type": g.gap_type,
                "statement": g.statement,
                "detail": g.detail,
                "evidence": g.evidence,
                "related_materials": g.related_materials,
                "actionability": g.actionability,
                "priority": g.priority,
                "source": g.source,
                "suggested_actions": g.suggested_actions,
                "subquery": g.subquery,
                "created_at": g.created_at.isoformat() if g.created_at else None,
            }
            for g in gaps
        ],
        "stats": stats,
    }


@app.get("/api/projects/{project_id}/conflicts")
def list_research_conflicts(project_id: str) -> dict:
    """获取文献冲突清单（交叉验证落库，供 Claim 冲突可视化与论文溯源）。

    每条冲突：冲突陈述 + 立场证据（support/refute 双方来源，可点击跳转
    论文页溯源）+ 处置建议 + 置信度 + 来源子问题。
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        conflicts = store.list_research_conflicts(limit=200)
        stats = store.conflict_stats()
    except Exception as e:  # noqa: BLE001
        return {"conflicts": [], "stats": {"total": 0, "papers": 0},
                "error": f"读取失败: {e}"}
    return {
        "conflicts": [
            {
                "conflict_id": c.conflict_id,
                "claim": c.claim,
                "sources": c.sources,
                "resolution": c.resolution,
                "confidence": c.confidence,
                "subquery": c.subquery,
                "created_at": c.created_at.isoformat() if c.created_at else None,
            }
            for c in conflicts
        ],
        "stats": stats,
    }


@app.post("/api/projects/{project_id}/materials/re-extract")
def re_extract_materials(project_id: str) -> dict:
    """覆盖度重抽：对「仅名称」材料（无性能无合成）做针对性二次抽取补全。

    仅名称材料是首轮抽取的召回遗漏（材料名出现但未抽到性能/合成）。
    本接口跨论文聚合包含材料名的摘要片段，专门 prompt 补全其性能与合成。
    异步执行（调用 LLM，可能耗时数十秒~数分钟），立即返回任务受理。
    """
    state = _require_project(project_id)
    if state.status == "running":
        raise HTTPException(status_code=409, detail="pipeline 正在运行中")

    def _worker() -> None:
        try:
            store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
            from core.llm import LLMRegistry
            from stages.research.agents import MaterialKnowledgeExtractionAgent
            registry = LLMRegistry.from_config(_CONFIG)
            agent = MaterialKnowledgeExtractionAgent("material_extraction")
            n_added = agent._re_extract_name_only(store, registry)
            mstats = store.material_stats()
            state.error = None
            state.summary = (
                f"覆盖度重抽完成：补全 {n_added} 条知识"
                f"（库内材料 {mstats['materials']} / 性能 {mstats['properties']} / "
                f"合成 {mstats['synthesis']}）"
            )
        except Exception as e:  # noqa: BLE001
            state.error = f"重抽失败: {type(e).__name__}: {e}"
            state.summary = f"覆盖度重抽异常: {e}"

    threading.Thread(target=_worker, daemon=True).start()
    return {
        "project_id": project_id,
        "message": "覆盖度重抽已启动（异步执行），完成后可刷新材料页查看补全结果",
    }


@app.get("/api/projects/{project_id}/evidence")
def list_evidence(project_id: str) -> dict:
    """获取检索证据链（审计轨迹：query → source → 命中 → paper 关联）。

    赛题手册明确要求文献调研可溯源：Sciverse 调用记录天然构成证据链。
    每条记录：触发子问题、数据源、命中证据标题、外部 ID（doc_id/arxiv_id）、
    证据分、片段摘要，以及是否关联到已入库论文。
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        entries = store.list_evidence(limit=500)
        stats = store.evidence_stats()
    except Exception as e:  # noqa: BLE001
        return {
            "entries": [],
            "stats": {"total": 0, "by_source": {}, "linked": 0},
            "error": f"读取失败: {e}",
        }
    return {"entries": entries, "stats": stats}


@app.get("/api/projects/{project_id}/unlinked-papers")
def list_unlinked_papers(project_id: str) -> dict:
    """列出未入库论文候选（检索命中但未关联入库的证据，按 external_id/title 去重）。

    检索阶段每条子问题按固定配额抓取（Sciverse 10 + arXiv 3 + S2 2），
    其中被相关性筛选或去重剔除的候选保留在证据链（paper_id 为空）。
    前端可展示为「未入库论文」，支持用户手动补录入库。

    每篇候选附带未入库原因（reason）：
    - score_rejected: 相关性打分 < 0.5 被 filter 剔除（Sciverse 用语义检索分，
      arXiv/S2 用 LLM 评估分）
    - dedup_merged:   与已入库论文重复（同 external_id / 同 title），去重时合并
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        entries = store.list_unlinked_evidence(limit=500)
        papers_all = store.list_papers()
    except Exception as e:  # noqa: BLE001
        return {"papers": [], "error": f"读取失败: {e}"}

    # 已入库论文的 arxiv_id / title 集合（用于判定去重合并）
    paper_ax: set[str] = set()
    paper_titles: set[str] = set()
    paper_id_by_ax: dict[str, str] = {}
    paper_id_by_title: dict[str, str] = {}
    for p in papers_all:
        ax = (p.arxiv_id or "").strip()
        tl = (p.title or "").strip().lower()
        if ax:
            paper_ax.add(ax)
            paper_id_by_ax.setdefault(ax, p.paper_id)
        if tl:
            paper_titles.add(tl)
            paper_id_by_title.setdefault(tl, p.paper_id)

    # 按 external_id（或 title）去重聚合：同一候选可能被多个子问题命中
    seen: dict[str, dict] = {}
    for e in entries:
        key = (e.get("external_id") or "").strip() or \
            (e.get("title") or "").strip().lower()
        if not key:
            continue
        if key not in seen:
            seen[key] = {
                "external_id": e.get("external_id", ""),
                "title": e.get("title", ""),
                "source": e.get("source", ""),
                "snippet": e.get("snippet", ""),
                "subquery": e.get("subquery", ""),
                "evidence_score": e.get("evidence_score", 0.0),
                "hit_count": 1,
                "log_ids": [e.get("log_id", "")],
                "reason": "",
                "reason_detail": "",
            }
        else:
            seen[key]["hit_count"] += 1
            if e.get("log_id"):
                seen[key]["log_ids"].append(e.get("log_id", ""))
            if not seen[key]["snippet"]:
                seen[key]["snippet"] = e.get("snippet", "")

    # 判定每篇候选的未入库原因
    for p in seen.values():
        ax = (p["external_id"] or "").strip()
        tl = (p["title"] or "").strip().lower()
        if (ax and ax in paper_ax) or (tl and tl in paper_titles):
            dup_pid = paper_id_by_ax.get(ax) or paper_id_by_title.get(tl, "")
            dup_title = next(
                (pp.title for pp in papers_all if pp.paper_id == dup_pid),
                "",
            )
            p["reason"] = "dedup_merged"
            p["reason_detail"] = (
                f"与已入库论文《{dup_title[:60]}》重复，去重时合并"
            )
        else:
            src = p["source"]
            scorer = "Sciverse 语义检索分" if src == "sciverse" else "LLM 相关性评估分"
            p["reason"] = "score_rejected"
            p["reason_detail"] = (
                f"相关性打分 < 0.5（{scorer}），被 filter 阶段剔除"
            )

    papers = sorted(
        seen.values(),
        key=lambda p: (p["evidence_score"] or 0.0),
        reverse=True,
    )
    return {
        "papers": papers,
        "total": len(papers),
        "filter_threshold": 0.5,
        "filter_note": (
            "入库量化标准：filter 阶段相关性打分 ≥ 0.5 才入库。"
            "Sciverse 候选复用语义检索证据分；arXiv/S2 候选由 LLM 按主题相关性评估。"
            "低于 0.5 的候选保留在未入库列表，可人工复核后手动补录。"
        ),
    }


@app.post("/api/projects/{project_id}/papers/import")
def import_unlinked_paper(project_id: str, body: dict) -> dict:
    """手动补录入库：将一条未关联证据候选正式入库为论文，并回填证据链关联。

    body:
      - external_id: 候选的外部 ID（sciverse doc_id / arxiv_id / s2 paperId）
      - title: 候选标题（external_id 缺失时的兜底定位）
      - snippet: 可选覆盖片段（否则用证据链中的 snippet 作为摘要）
    """
    _require_project(project_id)
    external_id = (body.get("external_id") or "").strip()
    title = (body.get("title") or "").strip()
    if not external_id and not title:
        raise HTTPException(status_code=400, detail="external_id 或 title 至少提供一个")

    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        entries = store.list_unlinked_evidence(limit=500)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"读取证据链失败: {e}")

    # 定位候选：优先 external_id 精确匹配，其次 title 完全一致
    cand = None
    for e in entries:
        eid = (e.get("external_id") or "").strip()
        etitle = (e.get("title") or "").strip().lower()
        if external_id and eid and eid == external_id:
            cand = e
            break
        if not external_id and title and etitle == title.lower():
            cand = e
            break
    if cand is None:
        # 候选不在未关联列表：可能已入库 → 返回已存在的论文信息
        if external_id:
            existing = store.find_paper_by_external_id(external_id)
            if existing is not None:
                return {
                    "paper_id": existing.paper_id,
                    "created": False,
                    "message": f"该候选已入库：{existing.title}",
                }
        raise HTTPException(
            status_code=404,
            detail="未找到该未入库候选（可能已被入库或已不在证据链中）",
        )

    # 重复入库检查：external_id 已存在则直接关联，不新建论文
    existing = None
    if (cand.get("external_id") or "").strip():
        existing = store.find_paper_by_external_id(cand["external_id"].strip())
    if existing is not None:
        linked_count = 0
        for e in entries:
            eid = (e.get("external_id") or "").strip()
            if external_id and eid and eid == external_id and e.get("log_id"):
                store.link_evidence_to_paper(e["log_id"], existing.paper_id)
                linked_count += 1
        return {
            "paper_id": existing.paper_id,
            "created": False,
            "message": f"该候选已存在同名论文，证据链已关联 {linked_count} 条到现有论文",
        }

    # 新建论文（用证据链中的元数据）
    paper_id = KnowledgeStore.new_id()
    snippet = body.get("snippet") or (cand.get("snippet") or "")
    from core.knowledge.schema import Paper
    paper = Paper(
        paper_id=paper_id,
        title=cand.get("title") or title or "Untitled",
        abstract=snippet or None,
        metadata={
            "source": cand.get("source", ""),
            "source_subquery": cand.get("subquery", ""),
            "doc_id": cand.get("external_id", ""),
            "evidence_score": cand.get("evidence_score", 0.0),
            "relevance_score": 0.0,
            "relevance_reason": "手动补录入库",
            "manually_imported": True,
        },
        source_stage="research",
    )
    # 入库论文 + 切分 chunk
    from core.tools.text_split import split_into_chunks
    store.save_paper(paper)
    abstract = snippet or ""
    chunks = split_into_chunks(abstract, max_tokens=500, overlap_tokens=50)
    from core.knowledge.schema import PaperChunk
    chunk_objs = [
        PaperChunk(
            chunk_id=KnowledgeStore.new_id(),
            paper_id=paper_id,
            chunk_index=i,
            text=c.text,
        )
        for i, c in enumerate(chunks)
    ]
    store.save_paper_chunks(chunk_objs)

    # 回填证据链关联：该候选可能被多个子问题命中，全部关联
    linked_count = 0
    for e in entries:
        eid = (e.get("external_id") or "").strip()
        etitle = (e.get("title") or "").strip().lower()
        match = bool(external_id and eid and eid == external_id) or \
            bool(not external_id and title and etitle == title.lower())
        if match and e.get("log_id"):
            store.link_evidence_to_paper(e["log_id"], paper_id)
            linked_count += 1

    return {
        "paper_id": paper_id,
        "created": True,
        "message": f"论文已入库：{paper.title}（关联 {linked_count} 条证据）",
    }


@app.get("/api/projects/{project_id}/papers/{paper_id}/pdf")
def download_paper_pdf(project_id: str, paper_id: str):
    """代理下载论文 PDF：后端拉取 PDF 存到本地，前端真正一键下载。

    流程：
    1. 读论文的 pdf_url（无则按 arxiv_id/doi 解析）
    2. 若 pdf_url 是 doi.org 兜底链接（会跳到 HTML 文章页），
       先尝试 Unpaywall/OpenAlex 找 OA PDF 直链，再解析落地页
       <meta name="citation_pdf_url"> 提取真实 PDF 地址
    3. 后端请求 PDF 内容，保存到 <project>/papers/<paper_id>.pdf
    4. 返回文件流（attachment 下载）
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        try:
            paper = store.get_paper(paper_id)
        except Exception:  # noqa: BLE001
            paper = None
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"读取论文失败: {e}") from e
    if paper is None:
        raise HTTPException(status_code=404, detail="论文不存在")

    # 1. 已有本地文件 → 直接返回
    if paper.pdf_path and os.path.exists(paper.pdf_path):
        return FileResponse(
            paper.pdf_path,
            media_type="application/pdf",
            filename=f"{paper.paper_id}.pdf",
        )

    # 2. 解析目标 PDF URL
    meta = paper.metadata or {}
    pdf_url = (meta.get("pdf_url") or "").strip()
    if not pdf_url and paper.arxiv_id:
        pdf_url = f"https://arxiv.org/pdf/{paper.arxiv_id}.pdf"
    if not pdf_url and paper.doi:
        pdf_url = f"https://doi.org/{paper.doi}"
    if not pdf_url:
        # 最后手段：OpenAlex 按标题反查 DOI + OA PDF（仅在用户主动点击下载时联网）
        try:
            from core.tools.doi_resolve import resolve_pdf_link
            meta_tmp = {
                "title": paper.title or "",
                "authors": paper.authors or [],
                "year": paper.year,
                "arxiv_id": paper.arxiv_id or "",
                "doi": paper.doi or "",
                "pdf_url": "",
            }
            pdf_url = resolve_pdf_link(meta_tmp)
            if pdf_url and meta_tmp.get("doi"):
                try:
                    from core.knowledge.schema import Paper as _Paper
                    store.save_paper(_Paper(
                        paper_id=paper.paper_id,
                        title=paper.title, authors=paper.authors, year=paper.year,
                        venue=paper.venue, arxiv_id=paper.arxiv_id,
                        abstract=paper.abstract, doi=meta_tmp["doi"], url=paper.url,
                        pdf_path=paper.pdf_path,
                        metadata={**meta, "pdf_url": pdf_url, "doi": meta_tmp["doi"]},
                        source_stage=paper.source_stage, created_at=paper.created_at,
                    ))
                except Exception:  # noqa: BLE001
                    pass
        except Exception:  # noqa: BLE001
            pdf_url = ""
    if not pdf_url:
        raise HTTPException(status_code=404, detail="该论文暂无可下载的 PDF 链接")

    # 3. 后端下载 PDF（带 HTML 落地页解析，处理 doi.org 兜底链接）
    import re
    import requests as _req

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/120.0 Safari/537.36"
        ),
        "Accept": "application/pdf,*/*;q=0.8,text/html;q=0.5",
        "Accept-Language": "en-US,en;q=0.9",
    }

    def _is_pdf(content: bytes) -> bool:
        return bool(content) and len(content) >= 1000 and b"%PDF" in content[:2000]

    def _download(url: str, timeout: int = 30) -> bytes:
        resp = _req.get(url, headers=headers, timeout=timeout, stream=True)
        resp.raise_for_status()
        return resp.content

    def _extract_pdf_from_html(html: bytes, base_url: str) -> str:
        """从出版商标注页解析真实 PDF 地址（citation_pdf_url / og:url 等）。"""
        text = html.decode("utf-8", errors="ignore")
        # 1. <meta name="citation_pdf_url" content="...">
        m = re.search(
            r'<meta[^>]+name=["\']citation_pdf_url["\'][^>]+content=["\']([^"\']+)["\']',
            text,
            re.I,
        )
        if not m:
            m = re.search(
                r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+name=["\']citation_pdf_url["\']',
                text,
                re.I,
            )
        if m:
            url = m.group(1).strip()
            if url.startswith("//"):
                url = "https:" + url
            return url
        # 2. <link rel="citation_pdf_url" href="...">
        m = re.search(
            r'<link[^>]+rel=["\']citation_pdf_url["\'][^>]+href=["\']([^"\']+)["\']',
            text,
            re.I,
        )
        if m:
            url = m.group(1).strip()
            if url.startswith("//"):
                url = "https:" + url
            return url
        # 3. og:url（兜底，通常是文章页本身）
        return ""

    content = b""
    tried = []
    # 候选 URL 列表：优先真实 PDF 直链，其次 OA 解析，再次落地页
    candidates: list[str] = []
    candidates.append(pdf_url)
    # 若当前是 doi.org 兜底，先尝试 Unpaywall/OpenAlex 找 OA PDF 直链（force 绕过缓存）
    if "doi.org" in pdf_url and paper.doi:
        try:
            oa = find_open_access_pdf(paper.doi, title=paper.title or "", force=True)
            if oa and oa not in candidates:
                candidates.append(oa)
        except Exception:  # noqa: BLE001
            pass

    for url in candidates:
        try:
            content = _download(url)
            if _is_pdf(content):
                break
            # 不是 PDF → 可能返回了 HTML 落地页，尝试解析 citation_pdf_url
            tried.append(url)
            pdf_href = _extract_pdf_from_html(content, url)
            if pdf_href:
                try:
                    content = _download(pdf_href)
                    if _is_pdf(content):
                        break
                except Exception:  # noqa: BLE001
                    pass
        except Exception as e:  # noqa: BLE001
            tried.append(f"{url} ({e})")
            content = b""

    if not _is_pdf(content):
        # 最后兜底：用标题再解析一次（OpenAlex 标题反查 DOI + OA PDF）
        try:
            from core.tools.doi_resolve import resolve_pdf_link
            meta_tmp = {
                "title": paper.title or "",
                "authors": paper.authors or [],
                "year": paper.year,
                "arxiv_id": paper.arxiv_id or "",
                "doi": paper.doi or "",
                "pdf_url": "",
            }
            final_url = resolve_pdf_link(meta_tmp)
            if final_url and final_url not in tried:
                content = _download(final_url)
                if not _is_pdf(content):
                    pdf_href = _extract_pdf_from_html(content, final_url)
                    if pdf_href:
                        content = _download(pdf_href)
        except Exception:  # noqa: BLE001
            content = b""

    if not _is_pdf(content):
        detail = (
            "该论文没有可直链下载的 PDF（doi.org 跳转的是 HTML 文章页，"
            "且 OpenAlex/Unpaywall 未找到开放获取版本）。"
            "通常意味着论文在付费墙内，请前往原文链接查看。"
        )
        raise HTTPException(status_code=502, detail=detail)

    # 4. 保存本地并返回
    try:
        proj_dir = _CONFIG.paths.project_dir(project_id)
        pdf_dir = os.path.join(proj_dir, "papers")
        os.makedirs(pdf_dir, exist_ok=True)
        local_path = os.path.join(pdf_dir, f"{paper_id}.pdf")
        with open(local_path, "wb") as f:
            f.write(content)
        # 回写 pdf_path（存 DB）
        try:
            from core.knowledge.schema import Paper as _Paper
            _paper = _Paper(
                paper_id=paper.paper_id,
                title=paper.title,
                authors=paper.authors,
                year=paper.year,
                venue=paper.venue,
                arxiv_id=paper.arxiv_id,
                abstract=paper.abstract,
                doi=paper.doi,
                url=paper.url,
                pdf_path=local_path,
                metadata=paper.metadata,
                source_stage=paper.source_stage,
                created_at=paper.created_at,
            )
            store.save_paper(_paper)
        except Exception:  # noqa: BLE001
            pass  # 回写失败不影响本次下载
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"保存 PDF 失败: {e}") from e

    return FileResponse(
        local_path,
        media_type="application/pdf",
        filename=f"{paper.paper_id}.pdf",
    )


@app.get("/api/projects/{project_id}/claims")
def list_claims(project_id: str, status: Optional[str] = None) -> dict:
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        claims = store.list_claims()
    except Exception as e:  # noqa: BLE001
        return {"claims": [], "error": f"读取失败: {e}"}
    items = []
    for c in claims:
        # 争议检测：claim 证据引用的论文与冲突 sources 有交集 → 争议中
        related_conflicts = []
        try:
            related_conflicts = store.conflicts_for_claim(c.evidence_refs)
        except Exception:  # noqa: BLE001
            related_conflicts = []
        item = {
            "claim_id": c.claim_id,
            "statement": c.statement,
            "status": c.status.value,
            "role": c.role,
            "evidence_count": len(c.evidence_refs),
            "evidence_refs": c.evidence_refs,
            "source_idea_id": c.source_idea_id,
            "created_at": c.created_at.isoformat() if c.created_at else None,
            "verified_at": c.verified_at.isoformat() if c.verified_at else None,
            # 冲突可视化：相关文献冲突（非空 → 该 Claim 处于争议中）
            "conflicts": [
                {
                    "conflict_id": x.conflict_id,
                    "claim": x.claim,
                    "sources": x.sources,
                    "resolution": x.resolution,
                    "confidence": x.confidence,
                    "subquery": x.subquery,
                }
                for x in related_conflicts
            ],
        }
        items.append(item)
    if status:
        items = [it for it in items if it["status"] == status]
    return {"claims": items}


@app.get("/api/projects/{project_id}/experiments")
def list_experiments(project_id: str) -> dict:
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        exps = store.list_experiments()
    except Exception as e:  # noqa: BLE001
        return {"experiments": [], "error": f"读取失败: {e}"}
    return {
        "experiments": [
            {
                "experiment_id": e.experiment_id,
                "name": e.name,
                "status": e.status.value,
                "verifies_claim_ids": e.verifies_claim_ids,
                "config": e.config,
                "result_summary": e.result_summary,
                "anomaly_notes": e.anomaly_notes,
                "started_at": e.started_at.isoformat() if e.started_at else None,
                "completed_at": e.completed_at.isoformat() if e.completed_at else None,
                "created_at": e.created_at.isoformat() if e.created_at else None,
            }
            for e in exps
        ]
    }


@app.post("/api/projects/{project_id}/notes")
def add_note(project_id: str, req: NoteRequest) -> dict:
    state = _require_project(project_id)
    text = (req.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text 不能为空")
    note = {
        "note_id": uuid.uuid4().hex,
        "text": text,
        "created_at": datetime.utcnow().isoformat(),
        "created_at": datetime.now().isoformat(),
    }
    state.notes.append(note)
    return {"note": note}


@app.get("/api/projects/{project_id}/notes")
def list_notes(project_id: str) -> dict:
    state = _require_project(project_id)
    return {"notes": list(state.notes)}


@app.post("/api/projects/{project_id}/human-response")
def submit_human_response(project_id: str, req: HumanResponseRequest) -> dict:
    state = _require_project(project_id)
    action = req.action or "continue"
    if action not in ("continue", "abort", "rollback"):
        raise HTTPException(status_code=400, detail="action 必须为 continue/abort/rollback")
    resp = HumanResponse(
        text=req.text,
        selected_option=req.selected_option,
        action=action,
    )
    # 检索偏好（年份/期刊）随人工响应透传给节点，并记录到项目状态
    if req.search_prefs:
        resp.context = {"search_prefs": req.search_prefs}
        state.search_prefs = dict(req.search_prefs)
    ok = _BRIDGE.submit(project_id, resp)
    if not ok:
        # 没有等待中的人工请求：可能已提交、已超时、被清除，或多实例分裂
        raise HTTPException(
            status_code=409,
            detail=(
                "当前没有等待中的人工请求：可能已提交、已超时或被清除。"
                "请刷新页面查看最新状态。"
            ),
        )
    return {"project_id": project_id, "submitted": True, "action": action}


# ===== 赛题对齐展示接口 =====


@app.get("/api/projects/{project_id}/research-report")
def get_research_report(project_id: str) -> dict:
    """文献调研报告（赛题基本任务核心产出）。

    返回 cross_validation_report，含 Research Gaps / 共识 / 冲突 / 整体置信度。
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        cv_report = store.get_kv("cross_validation_report")
    except Exception as e:  # noqa: BLE001
        return {"report": None, "error": f"读取失败: {e}"}
    if not cv_report:
        return {
            "report": None,
            "message": "尚未生成调研报告，请先运行 research 阶段（启动 Pipeline 或构效关系发现）",
        }
    # 补充论文计数
    try:
        paper_count = len(store.list_papers())
    except Exception:  # noqa: BLE001
        paper_count = 0
    return {
        "project_id": project_id,
        "report": cv_report,
        "paper_count": paper_count,
    }


@app.get("/api/projects/{project_id}/discovery-detail")
def get_discovery_detail(project_id: str) -> dict:
    """构效关系发现详细数据（路线 A 完整产出）。

    返回：
    - discovery_summary: 计数卡片
    - discovery_hypotheses: 假设列表
    - discovery_search_space: 搜索空间定义
    - discovery_relationships: 验证后的构效关系（含交叉验证结果）
    - discovery_search_trace: MCTS 搜索轨迹（前端可视化）
    - discovery_literature_points: 文献数据点（前端散点图）
    - discovery_report_content: Markdown 报告
    - materials_cross_validation_report: Materials Project 交叉验证报告
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        kv = store.list_kv()
    except Exception as e:  # noqa: BLE001
        return {"error": f"读取失败: {e}"}
    # 也补充 Claim 形式的发现（兼容旧项目）
    relationships_from_claims: list[dict] = []
    try:
        for c in store.list_claims():
            if c.source_stage == "discovery":
                relationships_from_claims.append({
                    "claim_id": c.claim_id,
                    "statement": c.statement,
                    "status": c.status.value,
                    "evidence_refs": c.evidence_refs,
                    "created_at": c.created_at.isoformat() if c.created_at else None,
                })
    except Exception:  # noqa: BLE001
        pass
    kv_relationships = kv.get("discovery_relationships", [])
    # 优先用 KV 完整版（含 cross_validation 字段），无则用 Claim 版
    relationships = kv_relationships if kv_relationships else relationships_from_claims
    return {
        "project_id": project_id,
        "discovery_summary": kv.get("discovery_summary", {}),
        "discovery_hypotheses": kv.get("discovery_hypotheses", []),
        "discovery_search_space": kv.get("discovery_search_space", {}),
        "discovery_relationships": relationships,
        "discovery_search_trace": kv.get("discovery_search_trace", {}),
        "discovery_literature_points": kv.get("discovery_literature_points", []),
        "discovery_report_content": kv.get("discovery_report_content", ""),
        "discovery_report_artifact_id": kv.get("discovery_report_artifact_id", ""),
        "materials_cross_validation_report": kv.get("materials_cross_validation_report", {}),
    }


@app.get("/api/projects/{project_id}/materials-cross-validation")
def get_materials_cross_validation(project_id: str) -> dict:
    """Materials Project 交叉验证报告（赛题路线 A 硬要求）。"""
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        report = store.get_kv("materials_cross_validation_report")
    except Exception as e:  # noqa: BLE001
        return {"report": None, "error": f"读取失败: {e}"}
    if not report:
        return {
            "report": None,
            "message": "尚未生成交叉验证报告，请先运行构效关系发现",
        }
    return {"project_id": project_id, "report": report}


@app.get("/api/projects/{project_id}/method-alignment")
def get_method_alignment(project_id: str) -> dict:
    """方法↔代码对齐（公式 LaTeX ↔ 实验代码关键词匹配）。

    从方法 Artifact 抽取 LaTeX 公式，与实验代码做关键词匹配，
    标注 mapped / partial / missing。
    """
    _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        # 方法 Artifact
        method_content = ""
        method_artifacts: list[dict] = []
        for art in store.list_artifact_versions("") if hasattr(store, "list_artifact_versions") else []:
            method_artifacts.append({"artifact_id": art.artifact_id, "title": art.title})
        # 简化版：扫描所有 METHOD_DOC 类型 Artifact
        try:
            with store._connect() as conn:  # type: ignore[attr-defined]
                rows = conn.execute(
                    "SELECT artifact_id, content FROM artifacts WHERE artifact_type = 'method_doc' ORDER BY created_at DESC LIMIT 5"
                ).fetchall()
                method_artifacts = [
                    {"artifact_id": r["artifact_id"], "content": r["content"][:5000]}
                    for r in rows
                ]
                if rows:
                    method_content = rows[0]["content"]
        except Exception:  # noqa: BLE001
            pass
        # 实验代码（取最近实验的 config，可能含代码片段）
        experiments = store.list_experiments()
        experiment_code_snippets: list[dict] = []
        for exp in experiments[:5]:
            experiment_code_snippets.append({
                "experiment_id": exp.experiment_id,
                "name": exp.name,
                "status": exp.status.value,
                "config_keys": list(exp.config.keys()) if exp.config else [],
                "result_summary": exp.result_summary,
            })
    except Exception as e:  # noqa: BLE001
        return {"error": f"读取失败: {e}"}

    # 简化版对齐：从 method_content 抽取 LaTeX 公式标记
    import re
    formulas: list[dict] = []
    if method_content:
        # 抽取 $$...$$ / \[...\] / equation 环境的公式
        formula_patterns = [
            (r"\$\$(.+?)\$\$", "display"),
            (r"\\\[(.+?)\\\]", "display"),
            (r"\\begin\{equation\}(.+?)\\end\{equation\}", "display"),
            (r"\$([^$\n]+?)\$", "inline"),
        ]
        seen = set()
        for pat, kind in formula_patterns:
            for m in re.finditer(pat, method_content, re.DOTALL):
                f = m.group(1).strip()
                if f and f not in seen and len(f) < 500:
                    seen.add(f)
                    # 关键词匹配：在实验代码片段中查找公式变量
                    keywords = re.findall(r"[a-zA-Z_][a-zA-Z0-9_]*", f)
                    keywords = [k for k in keywords if len(k) > 2 and k not in
                                ("frac", "sum", "int", "mathrm", "text", "left", "right",
                                 "begin", "end", "alpha", "beta", "gamma", "delta", "theta",
                                 "lambda", "mu", "sigma", "phi", "psi", "omega")]
                    matched_kw = []
                    for kw in keywords[:10]:
                        for exp in experiments:
                            cfg_str = str(exp.config or "")
                            if kw in cfg_str:
                                matched_kw.append(kw)
                                break
                    status = "mapped" if matched_kw else (
                        "partial" if keywords else "missing"
                    )
                    formulas.append({
                        "formula": f,
                        "kind": kind,
                        "keywords": keywords[:10],
                        "matched_keywords": matched_kw,
                        "alignment_status": status,
                    })
    return {
        "project_id": project_id,
        "method_artifacts": method_artifacts,
        "method_content_preview": method_content[:2000] if method_content else "",
        "experiment_code_snippets": experiment_code_snippets,
        "formulas": formulas,
        "alignment_summary": {
            "total_formulas": len(formulas),
            "mapped": sum(1 for f in formulas if f["alignment_status"] == "mapped"),
            "partial": sum(1 for f in formulas if f["alignment_status"] == "partial"),
            "missing": sum(1 for f in formulas if f["alignment_status"] == "missing"),
        },
    }


@app.get("/api/projects/{project_id}/dashboard")
def get_dashboard(project_id: str) -> dict:
    """首页 Dashboard 聚合数据（赛题对齐展示）。

    一次返回所有计数 + 调研报告摘要 + 发现摘要 + 交叉验证摘要，
    让前端 Dashboard 单次拉取即可渲染。
    """
    state = _require_project(project_id)
    try:
        store = KnowledgeStore(_CONFIG.paths.project_db(project_id))
        counts = {
            "papers": len(store.list_papers()),
            "ideas": len(store.list_ideas()),
            "claims": len(store.list_claims()),
            "experiments": len(store.list_experiments()),
            "discovery_claims": sum(
                1 for c in store.list_claims() if c.source_stage == "discovery"
            ),
        }
        cv_report = store.get_kv("cross_validation_report") or {}
        discovery_summary = store.get_kv("discovery_summary") or {}
        materials_cv = store.get_kv("materials_cross_validation_report") or {}
    except Exception as e:  # noqa: BLE001
        return {"error": f"读取失败: {e}"}

    # 方法对齐摘要（公式数）
    method_alignment_summary = {"total_formulas": 0}
    try:
        from core.artifacts import ArtifactManager as _AM
        from core.knowledge import ArtifactType
        am = _AM(_CONFIG.paths.project_db(project_id))
        # 统计 FORMULA + METHOD_DOC 类型产出
        total_formulas = 0
        for at in (ArtifactType.FORMULA, ArtifactType.METHOD_DOC):
            try:
                arts = am.list_artifacts(artifact_type=at)
                for a in arts:
                    content = a.content or {}
                    if isinstance(content, dict):
                        formulas = content.get("formula_code_map") or content.get("formulas") or []
                        total_formulas += len(formulas) if isinstance(formulas, list) else 0
                    else:
                        total_formulas += 1
            except Exception:  # noqa: BLE001
                pass
        method_alignment_summary["total_formulas"] = total_formulas
    except Exception:  # noqa: BLE001
        pass

    # 论文写作摘要
    writing_summary = {"artifact_count": 0}
    try:
        from core.artifacts import ArtifactManager as _AM
        from core.knowledge import ArtifactType
        am = _AM(_CONFIG.paths.project_db(project_id))
        art_count = 0
        for wt in (ArtifactType.PAPER_DRAFT, ArtifactType.REVIEW_NOTE):
            try:
                art_count += len(am.list_artifacts(artifact_type=wt))
            except Exception:  # noqa: BLE001
                pass
        writing_summary["artifact_count"] = art_count
    except Exception:  # noqa: BLE001
        pass

    # 阶段状态
    stage_statuses: dict[str, str] = {}
    current_stage = ""
    try:
        session = ProjectSession.load(project_id, _CONFIG.paths)
        current_stage = session.current_stage().value
        for stage in LifecycleStage.ordered():
            stage_statuses[stage.value] = session.status_of(stage).value
    except Exception:  # noqa: BLE001
        pass

    return {
        "project_id": project_id,
        "topic": state.topic,
        "status": state.status,
        "summary": state.summary,
        "current_stage": current_stage,
        "stage_statuses": stage_statuses,
        "counts": counts,
        "research_report_summary": {
            "gaps_count": len(cv_report.get("gaps", [])),
            "consensus_count": len(cv_report.get("consensus", [])),
            "conflicts_count": len(cv_report.get("conflicts", [])),
            "overall_confidence": cv_report.get("overall_confidence", 0),
        },
        "discovery_summary": discovery_summary,
        "materials_cv_summary": {
            "total_discoveries": materials_cv.get("total_discoveries", 0),
            "mp_validated": materials_cv.get("mp_validated", 0),
            "rule_validated": materials_cv.get("rule_validated", 0),
            "overall_confidence": materials_cv.get("overall_confidence", 0),
            "source": materials_cv.get("source", ""),
        },
        "method_alignment_summary": method_alignment_summary,
        "writing_summary": writing_summary,
        "pending_human": _BRIDGE.get_pending(project_id),
    }


# ===== 产出物下载 =====


@app.get("/api/projects/{project_id}/download/{artifact_type}")
def download_artifact(project_id: str, artifact_type: str, format: str = "md"):
    """下载关键产出物，支持 md / docx / pdf 三种格式。

    支持类型：
    - research-report：文献调研报告
    - discovery-report：构效关系发现报告
    - experiment-code：实验代码（仅 md/py 格式）
    - experiment-results：实验结果（含 metrics）
    - method-doc：方法文档
    - paper-draft：论文稿
    - claims-summary：Claim 汇总
    - ideas-summary：研究思路汇总
    - full-report：全流程综合报告
    """
    _require_project(project_id)
    store = KnowledgeStore(_CONFIG.paths.project_db(project_id))

    # 实验代码仅支持原文下载（py 格式）
    if artifact_type == "experiment-code":
        exp_dir = _CONFIG.paths.project_dir(project_id) / "experiments"
        code_path = exp_dir / "run_exp.py"
        if code_path.exists():
            code = code_path.read_text(encoding="utf-8")
        else:
            exps = store.list_experiments()
            code = f"# 无 run_exp.py，共有 {len(exps)} 条实验记录"
        return _make_download(code, "run_exp.py", "text/x-python", "md")

    # 其余类型先统一生成 Markdown 内容
    md_content = ""
    base_filename = ""

    if artifact_type == "research-report":
        report = store.get_kv("cross_validation_report") or {}
        md_content = _build_research_report_md(report)
        base_filename = "research_report"

    elif artifact_type == "discovery-report":
        content = store.get_kv("discovery_report_content") or ""
        if not content:
            summary = store.get_kv("discovery_summary") or {}
            content = f"# 构效关系发现报告\n\n{json.dumps(summary, ensure_ascii=False, indent=2)}"
        md_content = content
        base_filename = "discovery_report"

    elif artifact_type == "method-doc":
        try:
            from core.artifacts import ArtifactManager
            from core.knowledge import ArtifactType
            am = ArtifactManager(_CONFIG.paths.project_db(project_id))
            arts = am.list_artifacts(artifact_type=ArtifactType.METHOD_DOC)
            if arts:
                md_content = am.read_content(arts[0]) or str(arts[0].content or "")
            else:
                md_content = "# 无方法文档"
        except Exception:
            md_content = "# 无方法文档"
        base_filename = "method_doc"

    elif artifact_type == "paper-draft":
        try:
            from core.artifacts import ArtifactManager
            from core.knowledge import ArtifactType
            am = ArtifactManager(_CONFIG.paths.project_db(project_id))
            arts = am.list_artifacts(artifact_type=ArtifactType.PAPER_DRAFT)
            if arts:
                md_content = am.read_content(arts[0]) or str(arts[0].content or "")
            else:
                md_content = "# 无论文稿"
        except Exception:
            md_content = "# 无论文稿"
        base_filename = "paper_draft"

    elif artifact_type == "claims-summary":
        md_content = _build_claims_summary_md(store)
        base_filename = "claims_summary"

    elif artifact_type == "ideas-summary":
        md_content = _build_ideas_summary_md(store)
        base_filename = "ideas_summary"

    elif artifact_type == "experiment-results":
        md_content = _build_experiment_results_md(store, project_id)
        base_filename = "experiment_results"

    elif artifact_type == "full-report":
        md_content = _build_full_report_md(store, project_id)
        base_filename = "full_report"

    else:
        raise HTTPException(status_code=400, detail=f"不支持的下载类型: {artifact_type}")

    # 按格式转换
    fmt = (format or "md").lower()
    if fmt == "docx":
        return _md_to_docx_response(md_content, base_filename)
    elif fmt == "pdf":
        return _md_to_pdf_response(md_content, base_filename)
    else:
        return _make_download(md_content, f"{base_filename}.md", "text/markdown", "md")


def _build_research_report_md(report: dict) -> str:
    """把 cross_validation_report 转为 Markdown。"""
    md = "# 文献调研报告\n\n"
    md += f"**综合置信度**: {report.get('overall_confidence', 0):.2f}\n\n"

    gaps = report.get("gaps", [])
    if gaps:
        md += "## Research Gaps\n\n"
        for i, g in enumerate(gaps, 1):
            md += f"{i}. {g}\n"
        md += "\n"

    consensus = report.get("consensus", [])
    if consensus:
        md += "## 共识\n\n"
        for i, c in enumerate(consensus, 1):
            md += f"{i}. {c}\n"
        md += "\n"

    conflicts = report.get("conflicts", [])
    if conflicts:
        md += "## 冲突结论\n\n"
        for c in conflicts:
            if isinstance(c, dict):
                md += f"- **{c.get('topic', '?')}**: {c.get('description', '')}\n"
                if c.get("positions"):
                    for pos in c["positions"]:
                        md += f"  - {pos}\n"
            else:
                md += f"- {c}\n"
        md += "\n"

    return md


def _make_download(content: str, filename: str, media_type: str, fmt: str = "md") -> Response:
    """构造文件下载响应。"""
    return Response(
        content=content.encode("utf-8"),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_claims_summary_md(store: KnowledgeStore) -> str:
    """构建 Claim 汇总 Markdown。"""
    claims = store.list_claims()
    md = "# Claim 汇总\n\n"
    for c in claims:
        md += f"## {c.claim_id}\n\n"
        md += f"**陈述**: {c.statement}\n\n"
        md += f"**状态**: {c.status.value}\n\n"
        md += f"**角色**: {c.role}\n\n"
        if c.evidence_refs:
            md += "**证据**:\n"
            for ref in c.evidence_refs:
                md += f"- {ref.get('type', '?')}: {ref.get('id', '?')}"
                if ref.get("chunk_id"):
                    md += f" (chunk: {ref['chunk_id']})"
                md += "\n"
        md += "\n"
    return md


def _build_ideas_summary_md(store: KnowledgeStore) -> str:
    """构建研究思路汇总 Markdown。"""
    ideas = store.list_ideas()
    md = "# 研究思路汇总\n\n"
    md += f"共 {len(ideas)} 个思路\n\n---\n\n"
    for idea in ideas:
        md += f"## {idea.idea_id}\n\n"
        md += f"**状态**: {idea.status}\n\n"
        md += f"**思路**: {idea.text}\n\n"
        if idea.constraints:
            md += f"**约束**: {'; '.join(idea.constraints)}\n\n"
        if idea.source_paper_ids:
            md += f"**来源论文**: {', '.join(idea.source_paper_ids)}\n\n"
        if idea.validation_notes:
            md += f"**验证**: {json.dumps(idea.validation_notes, ensure_ascii=False)}\n\n"
        md += "---\n\n"
    return md


def _build_experiment_results_md(store: KnowledgeStore, project_id: str) -> str:
    """构建实验结果 Markdown（含 metrics）。"""
    experiments = store.list_experiments()
    md = "# 实验结果汇总\n\n"
    md += f"共 {len(experiments)} 个实验\n\n---\n\n"
    for exp in experiments:
        md += f"## {exp.name or exp.experiment_id}\n\n"
        md += f"**状态**: {exp.status.value if hasattr(exp.status, 'value') else exp.status}\n\n"
        if hasattr(exp, 'metrics') and exp.metrics:
            md += f"**Metrics**: {json.dumps(exp.metrics, ensure_ascii=False, indent=2)}\n\n"
        if exp.result_summary:
            md += f"**结果摘要**:\n```\n{exp.result_summary[:3000]}\n```\n\n"
        if exp.anomaly_notes:
            md += f"**异常**: {exp.anomaly_notes[:1000]}\n\n"
        if exp.verifies_claim_ids:
            md += f"**验证 Claim**: {', '.join(exp.verifies_claim_ids)}\n\n"
        md += "---\n\n"
    return md


def _build_full_report_md(store: KnowledgeStore, project_id: str) -> str:
    """构建全流程综合报告 Markdown。"""
    parts: list[str] = []

    # 标题
    topic = store.get_kv("research_topic") or "(未设置主题)"
    parts.append(f"# 科研项目全流程报告\n\n**研究主题**: {topic}\n\n**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n\n")

    # 1. 调研报告
    report = store.get_kv("cross_validation_report") or {}
    if report:
        parts.append("---\n\n# 第一部分：文献调研\n\n")
        parts.append(_build_research_report_md(report))

    # 2. 研究思路
    ideas = store.list_ideas()
    if ideas:
        parts.append("\n\n---\n\n# 第二部分：研究思路\n\n")
        parts.append(_build_ideas_summary_md(store))

    # 3. Claim 汇总
    claims = store.list_claims()
    if claims:
        parts.append("\n\n---\n\n# 第三部分：Claim 汇总\n\n")
        parts.append(_build_claims_summary_md(store))

    # 4. 方法文档
    try:
        from core.artifacts import ArtifactManager
        from core.knowledge import ArtifactType
        am = ArtifactManager(_CONFIG.paths.project_db(project_id))
        arts = am.list_artifacts(artifact_type=ArtifactType.METHOD_DOC)
        if arts:
            content = am.read_content(arts[0]) or str(arts[0].content or "")
            parts.append("\n\n---\n\n# 第四部分：方法设计\n\n")
            parts.append(content)
    except Exception:
        pass

    # 5. 实验结果
    experiments = store.list_experiments()
    if experiments:
        parts.append("\n\n---\n\n# 第五部分：实验结果\n\n")
        parts.append(_build_experiment_results_md(store, project_id))

    # 6. 论文稿
    try:
        from core.artifacts import ArtifactManager
        from core.knowledge import ArtifactType
        am = ArtifactManager(_CONFIG.paths.project_db(project_id))
        arts = am.list_artifacts(artifact_type=ArtifactType.PAPER_DRAFT)
        if arts:
            content = am.read_content(arts[0]) or str(arts[0].content or "")
            parts.append("\n\n---\n\n# 第六部分：论文稿\n\n")
            parts.append(content)
    except Exception:
        pass

    return "".join(parts) if parts else "# 全流程报告\n\n（暂无产出）"


# ===== 格式转换：Markdown → DOCX / PDF =====


def _md_to_docx_response(md_content: str, base_filename: str) -> Response:
    """将 Markdown 转为 DOCX 并返回下载响应。"""
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    in_code_block = False
    for line in md_content.split("\n"):
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]
            run = p.runs[0] if p.runs else p.add_run("")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # 标题
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped and stripped[0].isdigit() and ". " in stripped[:5]:
            idx = stripped.index(". ")
            doc.add_paragraph(stripped[idx + 2:], style="List Number")
        elif stripped:
            doc.add_paragraph(stripped)
        # 空行跳过

    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{base_filename}.docx"'},
    )


def _md_to_pdf_response(md_content: str, base_filename: str) -> Response:
    """将 Markdown 转为 PDF 并返回下载响应。

    使用 fpdf2 + Windows 系统中文字体（Microsoft YaHei / SimHei）。
    """
    import io
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 注册中文字体（Windows 系统字体）
    font_name = "chinese"
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    font_loaded = False
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdf.add_font(font_name, "", fp, uni=True)
                font_loaded = True
                break
            except Exception:
                continue
    if not font_loaded:
        # 兜底：用内置字体（不支持中文，但不会崩溃）
        font_name = "Helvetica"

    pdf.set_font(font_name, size=11)

    in_code_block = False
    for line in md_content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            pdf.set_font(font_name, size=8)
            # 截断超长行
            safe = stripped[:120] if len(stripped) > 120 else stripped
            pdf.multi_cell(0, 5, safe)
            pdf.set_font(font_name, size=11)
            continue

        if stripped.startswith("### "):
            pdf.set_font(font_name, size=13)
            pdf.ln(3)
            pdf.multi_cell(0, 7, stripped[4:])
            pdf.set_font(font_name, size=11)
        elif stripped.startswith("## "):
            pdf.set_font(font_name, size=15)
            pdf.ln(5)
            pdf.multi_cell(0, 8, stripped[3:])
            pdf.set_font(font_name, size=11)
        elif stripped.startswith("# "):
            pdf.set_font(font_name, size=18)
            pdf.ln(8)
            pdf.multi_cell(0, 10, stripped[2:])
            pdf.set_font(font_name, size=11)
        elif stripped.startswith("---"):
            pdf.ln(3)
            pdf.multi_cell(0, 5, "=" * 60)
            pdf.ln(2)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.multi_cell(0, 6, f"  • {stripped[2:]}")
        elif stripped:
            pdf.multi_cell(0, 6, stripped)
        else:
            pdf.ln(3)

    buf = io.BytesIO()
    pdf.output(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base_filename}.pdf"'},
    )


# ===== 文件上传 =====


@app.post("/api/projects/{project_id}/upload-paper")
async def upload_paper(
    project_id: str,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
):
    """上传 PDF/文本文献，入库为 Paper 实体。

    支持 .pdf / .txt / .md 文件。PDF 仅存储元信息（不做 OCR），
    txt/md 直接作为摘要与 chunk 素材。
    """
    _require_project(project_id)
    store = KnowledgeStore(_CONFIG.paths.project_db(project_id))

    # 读取文件内容
    raw = await file.read()
    filename = file.filename or "uploaded.txt"
    ext = Path(filename).suffix.lower()

    paper_id = f"paper_{uuid.uuid4().hex[:12]}"
    abstract = ""
    pdf_path = None
    parse_mode = "none"
    sections_count = 0
    figures_count = 0

    if ext == ".pdf":
        # PDF：保存到项目目录，调用 MinerU 解析结构化内容（赛题三·方向三推荐工具）
        upload_dir = _CONFIG.paths.project_dir(project_id) / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = upload_dir / filename
        pdf_path.write_bytes(raw)

        # MinerU 解析（无 API key 时降级到 pypdf fallback）
        from core.tools import parse_pdf_with_mineru, mineru_is_available
        if mineru_is_available():
            try:
                mineru_doc = parse_pdf_with_mineru(pdf_path)
                parse_mode = mineru_doc.mode
                sections_count = len(mineru_doc.sections)
                figures_count = len(mineru_doc.figures)
                # 优先使用 MinerU 抽取的 title
                extracted_title = mineru_doc.title if mineru_doc.title else None
                # 拼接前 2000 字作为 abstract
                abstract = mineru_doc.full_text[:2000] if mineru_doc.full_text else f"(MinerU {parse_mode} 解析: {filename})"
                # 把 MinerU 解析的结构化数据存入 KV（前端可视化与 discovery 阶段使用）
                store.save_kv(f"mineru_{paper_id}", {
                    "title": mineru_doc.title,
                    "sections_count": sections_count,
                    "figures_count": figures_count,
                    "tables_count": len(mineru_doc.tables),
                    "equations_count": len(mineru_doc.equations),
                    "references_count": len(mineru_doc.references),
                    "mode": parse_mode,
                    "sections": [
                        {"heading": s.heading, "level": s.level, "page": s.page}
                        for s in mineru_doc.sections[:30]
                    ],
                })
                if extracted_title and not title:
                    title = extracted_title
            except Exception as e:
                parse_mode = "error"
                abstract = f"(MinerU 解析失败: {e}，需手动提取文本)"
        else:
            parse_mode = "fallback"
            abstract = f"(PDF 文件已上传: {filename}，未启用 MinerU 解析，需手动提取文本)"

    elif ext in (".txt", ".md"):
        parse_mode = "text"
        abstract = raw.decode("utf-8", errors="replace")[:5000]
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}")

    from core.knowledge import Paper
    paper = Paper(
        paper_id=paper_id,
        title=title or filename,
        authors=[],
        abstract=abstract,
        url=None,
        pdf_path=str(pdf_path) if pdf_path else None,
        source_stage="upload",
    )
    store.save_paper(paper)

    # txt/md 内容切分为 chunk（PDF 由 MinerU 解析后也会按 section 切分）
    chunk_count = 0
    if ext in (".txt", ".md") and abstract:
        from core.tools import split_into_chunks
        from core.knowledge import PaperChunk
        text_chunks = split_into_chunks(abstract, max_tokens=500, overlap_tokens=50)
        paper_chunks = [
            PaperChunk(
                chunk_id=f"{paper_id}_c{tc.index}",
                paper_id=paper_id,
                chunk_index=tc.index,
                text=tc.text,
            )
            for tc in text_chunks
        ]
        if paper_chunks:
            store.save_paper_chunks(paper_chunks)
            chunk_count = len(paper_chunks)
    elif ext == ".pdf" and parse_mode != "none" and parse_mode != "error":
        # MinerU 解析后按 section 切分 chunk（结构化优于字符切分）
        try:
            from core.tools import parse_pdf_with_mineru
            from core.knowledge import PaperChunk
            mineru_doc = parse_pdf_with_mineru(pdf_path)
            paper_chunks = []
            for i, sec in enumerate(mineru_doc.sections):
                if not sec.text or len(sec.text.strip()) < 20:
                    continue
                paper_chunks.append(PaperChunk(
                    chunk_id=f"{paper_id}_s{i}",
                    paper_id=paper_id,
                    chunk_index=i,
                    text=sec.text[:1500],
                ))
                if len(paper_chunks) >= 30:  # 上限 30 section chunks
                    break
            if paper_chunks:
                store.save_paper_chunks(paper_chunks)
                chunk_count = len(paper_chunks)
        except Exception:
            pass

    return {
        "paper_id": paper_id,
        "title": paper.title,
        "filename": filename,
        "chunks": chunk_count,
        "parse_mode": parse_mode,
        "sections_count": sections_count,
        "figures_count": figures_count,
        "message": "文献上传成功",
    }


@app.post("/api/projects/{project_id}/upload-topic")
async def upload_topic(
    project_id: str,
    file: UploadFile = File(...),
):
    """上传主题描述文件（.txt/.md），覆盖项目主题。"""
    state = _require_project(project_id)
    raw = await file.read()
    text = raw.decode("utf-8", errors="replace").strip()
    if not text:
        raise HTTPException(status_code=400, detail="文件内容为空")
    state.topic = text
    return {"project_id": project_id, "topic": text, "message": "主题已更新"}


# ===== 辅助 =====


def _require_project(project_id: str) -> ProjectState:
    with _LOCK:
        state = _PROJECTS.get(project_id)
    if state is None:
        raise HTTPException(
            status_code=404,
            detail=(
                f"项目不存在: {project_id}"
                "（若刚重启服务，内存项目列表已丢失，请重新创建项目；"
                "若访问了错误端口，请确认 8001 单实例）"
            ),
        )
        raise HTTPException(status_code=404, detail=f"项目不存在: {project_id}")
    return state


# ===== 静态资源 =====


@app.get("/")
def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/static/{file_path:path}")
def static_files(file_path: str) -> FileResponse:
    """静态资源（强制 no-cache，避免前端更新后用户仍看到旧版缓存导致样式错乱）。"""
    target = (STATIC_DIR / file_path).resolve()
    if not str(target).startswith(str(STATIC_DIR.resolve())) or not target.is_file():
        raise HTTPException(status_code=404, detail="Not Found")
    return FileResponse(
        target,
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )
# 挂载静态目录（/static/...）
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("SRA_WEB_PORT", "8001"))
    uvicorn.run("web.api:app", host="0.0.0.0", port=port, reload=False)
    uvicorn.run("web.api:app", host="0.0.0.0", port=8000, reload=False)
